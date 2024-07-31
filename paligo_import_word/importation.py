from paligo_requests.paligo_requests import Paligo_request
from docx import Document
from lxml import etree
from bs4 import BeautifulSoup as Soup
import os
import shutil
from paligo_import_word.classifier_bot.classifier_run import run_supervised_learning

class Importation:
    def __init__(self, paligo_output_folder_id: int, importation_document_path: str):
        self.importation_document_path = importation_document_path
        self.paligo_output_folder_id = paligo_output_folder_id
        self.paligo_r = Paligo_request("prod")
        self.analyse_input()
    
    def analyse_input(self):    
        extention = self.importation_document_path.split('.')[-1]
        #print(extention)
        if extention == "docx":
            self.import_documents_docx()
            
     
    def import_documents_docx(self):
        image_folder = self.extract_images("db\\import_documents\\images")
        xml = self.docx_to_xml(image_folder)
        
        xml_to_Docbook = self.xml_to_docBook(xml)
    
    def xml_to_docBook(self, xml_str: str):
        classified_text = []
        content_soup = Soup(xml_str, "xml")
        #print(content_soup)
        for para in content_soup.find_all("paragraph"):
            #print(para)
            data = run_supervised_learning(para)
            print(data)
            classified_text.append(data)
        
    def docx_to_xml(self, images_folder: str):
        docx_file = self.importation_document_path
    # Load the .docx document
        doc = Document(docx_file)
        print(doc)
        # Create the root element for the XML tree
        root = etree.Element("document")
        
        # Iterate through paragraphs in the documen
        
        
        # Iterate through paragraphs in the document
        for paragraph in doc.paragraphs:
            # Create a new XML element for each paragraph
            if paragraph.style.name == "Heading 1":
                para_elem = etree.SubElement(root, "title")
                para_elem.text = paragraph.text
            else:
                para_elem = etree.SubElement(root, "paragraph")
                para_elem.text = paragraph.text
        
        # Iterate through images in the folder
        image_files = os.listdir(images_folder)
        for image_file in image_files:
            # Create a new XML element for each image
            image_elem = etree.SubElement(root, "image")
            image_elem.set("src", os.path.join(images_folder, image_file))
        
        # Convert the XML tree to a string
        xml_str = etree.tostring(root, pretty_print=True, encoding="unicode")

    
        
        # Iterate through tables in the document
        for table in doc.tables:
            # Create a new XML element for each table
            table_elem = etree.SubElement(root, "table")
            for row in table.rows:
                row_elem = etree.SubElement(table_elem, "row")
                for cell in row.cells:
                    cell_elem = etree.SubElement(row_elem, "cell")
                    cell_elem.text = cell.text
        
        # Convert the XML tree to a string
        xml_str = etree.tostring(root, pretty_print=True, encoding="unicode")
        
        return xml_str
    def extract_images(self, output_folder):
    # Create the output folder if it doesn't exist
        docx_file = self.importation_document_path
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
        
        # Load the .docx document
        doc = Document(docx_file)
        
        # Iterate through the inline shapes (images) in the document
        for i, shape in enumerate(doc.inline_shapes):
            # Extract the image data
            try:
                image_data = shape.image.blob
                image_filename = f"image_{i}.png"  # Create a unique filename for each image
                
                # Save the image to the output folder
                image_path = os.path.join(output_folder, image_filename)
                with open(image_path, "wb") as f:
                    f.write(image_data)
            except Exception as e:
                print(e)
        
        return output_folder
            
if __name__ == '__main__':
    importation = Importation(paligo_output_folder_id=42658038, importation_document_path= "db\\import_documents\\Chap_1_SADR_VF_CH FINAL.docx")